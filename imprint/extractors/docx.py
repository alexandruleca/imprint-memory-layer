"""Microsoft Word (.docx) extractor via python-docx.

Legacy .doc (binary Word 97-2003) is not supported — recommend converting
with LibreOffice or antiword before ingesting.
"""

from __future__ import annotations

import io
import os

from . import ExtractedDoc, ExtractorUnavailable, ExtractionError, register_ext, register_mime


def _extract_from_fileobj(fileobj) -> tuple[str, dict]:
    try:
        import docx  # type: ignore  # python-docx package
    except ImportError as e:
        raise ExtractorUnavailable("python-docx not installed — `pip install python-docx`") from e

    doc = docx.Document(fileobj)
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]

    # Tables — flatten row-by-row so content in them isn't lost.
    table_text = []
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                table_text.append(" | ".join(cells))

    text = "\n\n".join(paras + table_text).strip()

    meta: dict = {}
    try:
        props = doc.core_properties
        for src, dst in [
            ("title", "title"), ("author", "author"),
            ("subject", "subject"), ("keywords", "keywords"),
        ]:
            v = getattr(props, src, None)
            if v:
                meta[dst] = str(v)
    except Exception:
        pass
    return text, meta


def _extract_bytes(data: bytes, source_url: str = "") -> ExtractedDoc:
    try:
        text, meta = _extract_from_fileobj(io.BytesIO(data))
    except ExtractorUnavailable:
        raise
    except Exception as e:
        raise ExtractionError(f"docx parse failed: {e}") from e
    return ExtractedDoc(
        text=text,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        metadata=meta,
        chunk_mode="prose",
    )


def extract(path: str) -> ExtractedDoc:
    with open(path, "rb") as f:
        doc = _extract_bytes(f.read())
    doc.metadata.setdefault("filename", os.path.basename(path))
    return doc


register_ext(".docx", extract)
register_mime(
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    _extract_bytes,
)
